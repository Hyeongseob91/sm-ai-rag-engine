"""
File Parsers - 다양한 파일 형식 파싱.

PDF, DOCX, XLSX, TXT, JSON 등 다양한 파일 형식을 파싱하여 RawDocument 형태로 변환합니다.
"""
import json
from abc import ABC, abstractmethod
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List

import pdfplumber
from docx import Document
from openpyxl import load_workbook

from ...schemas.preprocessing import RawDocument


class BaseParser(ABC):
    """모든 파서의 기본 클래스.

    추상 메서드를 통해 각 파일 형식별 파서가 구현해야 할 인터페이스를 정의합니다.
    """

    @property
    @abstractmethod
    def supported_extensions(self) -> List[str]:
        """지원하는 파일 확장자 목록"""
        pass

    @abstractmethod
    def parse(self, file_path: str) -> RawDocument:
        """파일을 파싱하여 RawDocument 반환"""
        pass

    def can_you_parse(self, file_path: str) -> bool:
        """파일 파싱 가능 여부 확인"""
        ext = Path(file_path).suffix.lower().lstrip(".")
        return ext in self.supported_extensions

    def _get_file_info(self, file_path: str) -> Dict[str, Any]:
        """파일 기본 정보 추출"""
        path = Path(file_path)
        stat = path.stat()
        return {
            "file_name": path.name,
            "file_type": path.suffix.lower().lstrip("."),
            "file_size": stat.st_size,
            "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
            "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
        }


class PDFParser(BaseParser):
    """PDF 파일 파서 (pdfplumber).

    pdfplumber를 사용하여 PDF 파일의 텍스트와 테이블을 추출하고 페이지별로 구조화합니다.
    """

    @property
    def supported_extensions(self) -> List[str]:
        return ["pdf"]

    def parse(self, file_path: str) -> RawDocument:
        file_info = self._get_file_info(file_path)
        pages = []

        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                pages.append(text)

                # 테이블 추출 (있는 경우)
                tables = page.extract_tables()
                for table in tables:
                    if table:
                        table_text = self._table_to_text(table)
                        if table_text not in text:
                            pages[-1] += f"\n\n[Table]\n{table_text}"

            file_info["page_count"] = len(pdf.pages)
            file_info["pdf_metadata"] = pdf.metadata or {}

        content = "\n\n".join(pages)

        return RawDocument(
            content=content,
            source=str(Path(file_path).absolute()),
            file_type="pdf",
            file_name=file_info["file_name"],
            metadata=file_info,
            pages=pages,
        )

    def _table_to_text(self, table: List[List]) -> str:
        """테이블을 텍스트로 변환"""
        rows = []
        for row in table:
            cells = [str(cell or "").strip() for cell in row]
            rows.append(" | ".join(cells))
        return "\n".join(rows)


class DOCXParser(BaseParser):
    """DOCX 파일 파서 (python-docx).

    python-docx를 사용하여 Word 문서의 문단과 테이블을 추출하고 문서 속성 정보를 포함합니다.
    """

    @property
    def supported_extensions(self) -> List[str]:
        return ["docx"]

    def parse(self, file_path: str) -> RawDocument:
        file_info = self._get_file_info(file_path)
        doc = Document(file_path)

        paragraphs = []

        # 문단 추출
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # 테이블 추출
        for table in doc.tables:
            table_text = self._extract_table(table)
            if table_text:
                paragraphs.append(f"[Table]\n{table_text}")

        content = "\n\n".join(paragraphs)

        # 문서 속성 추출
        props = doc.core_properties
        file_info["title"] = props.title or ""
        file_info["author"] = props.author or ""
        file_info["paragraph_count"] = len(doc.paragraphs)
        file_info["table_count"] = len(doc.tables)

        return RawDocument(
            content=content,
            source=str(Path(file_path).absolute()),
            file_type="docx",
            file_name=file_info["file_name"],
            metadata=file_info,
            pages=paragraphs,  # 문단을 페이지처럼 취급
        )

    def _extract_table(self, table) -> str:
        """테이블을 텍스트로 변환"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        return "\n".join(rows)


class XLSXParser(BaseParser):
    """XLSX 파일 파서 (openpyxl).

    openpyxl을 사용하여 엑셀 파일의 모든 시트를 파싱하고 시트별로 구조화된 데이터를 제공합니다.
    """

    @property
    def supported_extensions(self) -> List[str]:
        return ["xlsx", "xls"]

    def parse(self, file_path: str) -> RawDocument:
        file_info = self._get_file_info(file_path)
        wb = load_workbook(file_path, read_only=True, data_only=True)

        sheets = {}
        all_content = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheet_content = self._extract_sheet(ws)

            if sheet_content.strip():
                sheets[sheet_name] = sheet_content
                all_content.append(f"[Sheet: {sheet_name}]\n{sheet_content}")

        content = "\n\n".join(all_content)

        file_info["sheet_count"] = len(wb.sheetnames)
        file_info["sheet_names"] = wb.sheetnames

        wb.close()

        return RawDocument(
            content=content,
            source=str(Path(file_path).absolute()),
            file_type="xlsx",
            file_name=file_info["file_name"],
            metadata=file_info,
            sheets=sheets,
        )

    def _extract_sheet(self, ws) -> str:
        """시트 내용을 텍스트로 변환"""
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(cell) if cell is not None else "" for cell in row]
            if any(cells):  # 빈 행 제외
                rows.append(" | ".join(cells))
        return "\n".join(rows)


class TXTParser(BaseParser):
    """TXT 파일 파서.

    다양한 인코딩(UTF-8, CP949, EUC-KR 등)을 자동 감지하여 텍스트 파일을 안전하게 읽습니다.
    """

    @property
    def supported_extensions(self) -> List[str]:
        return ["txt", "md", "rst"]

    def parse(self, file_path: str) -> RawDocument:
        file_info = self._get_file_info(file_path)

        # 인코딩 자동 감지 시도
        encodings = ["utf-8", "cp949", "euc-kr", "latin-1"]
        content = None

        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    content = f.read()
                file_info["encoding"] = encoding
                break
            except UnicodeDecodeError:
                continue

        if content is None:
            raise ValueError(f"파일 인코딩을 감지할 수 없습니다: {file_path}")

        file_info["line_count"] = content.count("\n") + 1
        file_info["char_count"] = len(content)

        return RawDocument(
            content=content,
            source=str(Path(file_path).absolute()),
            file_type=file_info["file_type"],
            file_name=file_info["file_name"],
            metadata=file_info,
        )


class JSONParser(BaseParser):
    """JSON 파일 파서.

    JSON 데이터를 계층적 구조를 유지한 채 읽기 쉬운 텍스트 형식으로 변환합니다.
    """

    @property
    def supported_extensions(self) -> List[str]:
        return ["json"]

    def parse(self, file_path: str) -> RawDocument:
        file_info = self._get_file_info(file_path)

        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        # JSON을 읽기 쉬운 텍스트로 변환
        content = self._json_to_text(data)

        file_info["json_type"] = type(data).__name__
        if isinstance(data, dict):
            file_info["top_level_keys"] = list(data.keys())
        elif isinstance(data, list):
            file_info["item_count"] = len(data)

        return RawDocument(
            content=content,
            source=str(Path(file_path).absolute()),
            file_type="json",
            file_name=file_info["file_name"],
            metadata=file_info,
        )

    def _json_to_text(self, data: Any, prefix: str = "") -> str:
        """JSON 데이터를 텍스트로 변환 (계층적 구조 유지)"""
        lines = []

        if isinstance(data, dict):
            for key, value in data.items():
                if isinstance(value, (dict, list)):
                    lines.append(f"{prefix}{key}:")
                    lines.append(self._json_to_text(value, prefix + "  "))
                else:
                    lines.append(f"{prefix}{key}: {value}")
        elif isinstance(data, list):
            for i, item in enumerate(data):
                if isinstance(item, (dict, list)):
                    lines.append(f"{prefix}[{i}]:")
                    lines.append(self._json_to_text(item, prefix + "  "))
                else:
                    lines.append(f"{prefix}- {item}")
        else:
            lines.append(f"{prefix}{data}")

        return "\n".join(lines)


class UnifiedFileParser:
    """통합 파일 파서 - 파일 확장자에 따라 적절한 파서 선택.

    등록된 모든 파서를 순회하여 파일 확장자에 맞는 파서를 자동으로 선택하고 실행합니다.
    """

    def __init__(self):
        self._parsers: List[BaseParser] = [
            PDFParser(),
            DOCXParser(),
            XLSXParser(),
            TXTParser(),
            JSONParser(),
        ]

    def get_supported_extensions(self) -> List[str]:
        """지원하는 모든 파일 확장자 반환"""
        extensions = []
        for parser in self._parsers:
            extensions.extend(parser.supported_extensions)
        return extensions

    def parse(self, file_path: str) -> RawDocument:
        """파일을 파싱하여 RawDocument 반환"""
        for parser in self._parsers:
            if parser.can_you_parse(file_path):
                return parser.parse(file_path)

        ext = Path(file_path).suffix
        raise ValueError(f"지원하지 않는 파일 형식입니다: {ext}")

    def parse_directory(self, dir_path: str, recursive: bool = False) -> List[RawDocument]:
        """디렉토리 내 모든 지원 파일 파싱"""
        documents = []
        path = Path(dir_path)
        pattern = "**/*" if recursive else "*"

        for file_path in path.glob(pattern):
            if file_path.is_file():
                ext = file_path.suffix.lower().lstrip(".")
                if ext in self.get_supported_extensions():
                    try:
                        doc = self.parse(str(file_path))
                        documents.append(doc)
                    except Exception as e:
                        print(f"파싱 실패: {file_path.name} - {e}")

        return documents
